#! /usr/bin/python3

# Brief WordX main program
# Author: Ashad Mohamed (aka. mashad)
# **********************************
# This the convert json to docx implementation for this project, contain all general function
# to implement docx manipulation.

import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def apply_run_style(run, style):
    if 'bold' in style and style['bold']:
        run.bold = True
    if 'italic' in style and style['italic']:
        run.italic = True
    if 'underline' in style and style['underline']:
        run.underline = True
    # if 'color' in style and style['color']:
    #     run.font.color.rgb = style['color']
    if 'size' in style and style['size']:
        run.font.size = Pt(style['size'])
    if 'name' in style and style['name']:
        run.font.name = style['name']

def apply_paragraph_alignment(para, alignment):
    alignments = {
        'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
        'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
        'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
        'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    }
    para.alignment = alignments.get(alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)  # Default to 'left'

def generate_word_from_json(json_file, output_docx):
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()
    
    for element in data:
        if element['type'] == 'paragraph':
            para = doc.add_paragraph()
            apply_paragraph_alignment(para, element.get('alignment', 'left'))  # Apply alignment
            for run_data in element['runs']:
                run = para.add_run(run_data['text'])
                apply_run_style(run, run_data['style'])
        elif element['type'] == 'header':
            level = element['level']
            para = doc.add_heading(level=level)
            para.add_run(element['text'])

    doc.save(output_docx)
    print(f"Successfully saved content from {json_file} to {output_docx}.")

def convertJsonToDocx(json, output):
    generate_word_from_json(json, output)
