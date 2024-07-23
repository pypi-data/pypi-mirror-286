import json
import docx2python
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# TO-DO: Function Add Comment
# Add a comment
def addcomment(file_path: str, text_to_comment: str, comment_text: str, author: str, timestamp: str) -> None:
    doc = Document(file_path)
    comment_id = 0

    # Create comments part if it does not exist
    comments_part = doc.part._element.xpath('/w:document/w:comments')
    if not comments_part:
        comments_part = OxmlElement('w:comments')
        doc.part._element.append(comments_part)
    else:
        comments_part = comments_part[0]

    new_comment = OxmlElement('w:comment')
    new_comment.set(qn('w:id'), str(comment_id))
    new_comment.set(qn('w:author'), author)
    new_comment.set(qn('w:date'), timestamp)
    new_comment.append(OxmlElement('w:p'))
    new_comment[0].append(OxmlElement('w:r'))
    new_comment[0][0].append(OxmlElement('w:t'))
    new_comment[0][0][0].text = comment_text
    comments_part.append(new_comment)

    for paragraph in doc.paragraphs:
        if text_to_comment in paragraph.text:
            start = paragraph.text.find(text_to_comment)
            end = start + len(text_to_comment)
            run = paragraph.add_run()
            comment_range_start = OxmlElement('w:commentRangeStart')
            comment_range_start.set(qn('w:id'), str(comment_id))
            run._r.insert(0, comment_range_start)
            run = paragraph.add_run()
            comment_range_end = OxmlElement('w:commentRangeEnd')
            comment_range_end.set(qn('w:id'), str(comment_id))
            run._r.append(comment_range_end)
            run = paragraph.add_run()
            comment_reference = OxmlElement('w:commentReference')
            comment_reference.set(qn('w:id'), str(comment_id))
            run._r.append(comment_reference)

    doc.save(file_path)

